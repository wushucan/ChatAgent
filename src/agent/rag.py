"""RAG engine — document loading, chunking, embedding, and retrieval.

Uses Chroma for vector storage and ``FastEmbedEmbeddings`` with
``BAAI/bge-small-zh-v1.5`` for local text embedding.
"""

from __future__ import annotations

import json
import os
import threading
import uuid
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from langchain_community.embeddings import FastEmbedEmbeddings
from langchain_community.vectorstores import Chroma
from langchain_text_splitters import RecursiveCharacterTextSplitter

# 国内 HuggingFace 镜像，首次下载模型用
os.environ.setdefault("HF_ENDPOINT", "https://hf-mirror.com")


# ── 文档解析器 ────────────────────────────────────────────────────


def _parse_txt(file_path: str) -> str:
    with open(file_path, "r", encoding="utf-8", errors="replace") as f:
        return f.read()


def _parse_md(file_path: str) -> str:
    return _parse_txt(file_path)


def _parse_pdf(file_path: str) -> str:
    import fitz

    doc = fitz.open(file_path)
    pages = []
    for page in doc:
        text = page.get_text()
        if text.strip():
            pages.append(text)
    doc.close()
    return "\n\n".join(pages)


def _parse_docx(file_path: str) -> str:
    from docx import Document
    import zipfile
    from lxml import etree

    doc = Document(file_path)
    parts = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)

    # 如果 python-docx 提取不到文本（如文本框/形状排版），回退到 XML 直接解析
    if not parts:
        with zipfile.ZipFile(file_path) as z:
            xml = z.read("word/document.xml")
            root = etree.fromstring(xml)
            ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
            texts = root.findall(".//w:t", ns)
            parts = [t.text for t in texts if t.text and t.text.strip()]

    return "\n".join(parts)


_PARSERS = {
    ".txt": _parse_txt,
    ".md": _parse_md,
    ".pdf": _parse_pdf,
    ".docx": _parse_docx,
}


# ── RAG 引擎 ───────────────────────────────────────────────────────


class RAGEngine:
    """本地文档知识库，支持 PDF/TXT/MD/DOCX 的索引与检索。"""

    def __init__(self, data_dir: str):
        self._data_dir = Path(data_dir)
        self._index_path = self._data_dir / "index.json"
        self._persist_dir = str(self._data_dir / ".rag_index")

        self._index_path.parent.mkdir(parents=True, exist_ok=True)
        self._embeddings: FastEmbedEmbeddings | None = None
        self._vectorstore: Chroma | None = None
        self._model_ready = False
        self._model_error: str | None = None
        self._splitter = RecursiveCharacterTextSplitter(
            chunk_size=500,
            chunk_overlap=50,
            length_function=len,
            separators=["\n\n", "\n", "。", ".", " ", ""],
        )

    # ── 懒初始化 ──────────────────────────────────────────────────

    def _ensure_embeddings(self):
        if self._embeddings is None:
            self._embeddings = FastEmbedEmbeddings(
                model_name="BAAI/bge-small-zh-v1.5",
            )

    @property
    def model_ready(self) -> bool:
        return self._model_ready

    @property
    def model_error(self) -> str | None:
        return self._model_error

    def warmup(self):
        """在后台线程预加载 embedding 模型，避免首次上传阻塞 UI。"""
        import sys
        try:
            self._ensure_embeddings()
            self._model_ready = True
        except Exception as e:
            print(f"[rag] warmup error: {e}", file=sys.stderr)
            self._model_error = str(e)

    def _ensure_vectorstore(self):
        self._ensure_embeddings()
        if self._vectorstore is None:
            self._vectorstore = Chroma(
                persist_directory=self._persist_dir,
                embedding_function=self._embeddings,
            )

    # ── 文档管理 ──────────────────────────────────────────────────

    def add_document(self, file_path: str) -> dict[str, Any]:
        """解析文件 → 分块 → 嵌入 → 存入 Chroma。

        Returns:
            成功时: ``{"id": str, "filename": str, "chunks": int, "created_at": str}``
            失败时: ``{"error": str}``
        """
        src = Path(file_path)
        ext = src.suffix.lower()
        parser = _PARSERS.get(ext)
        if parser is None:
            return {"error": f"不支持的文件格式: {ext}"}

        raw_text = parser(file_path)
        if not raw_text.strip():
            return {"error": "无法提取文本内容，文件可能为空或是扫描件"}

        chunks = self._splitter.split_text(raw_text)
        if not chunks:
            return {"error": "文件内容过短，无法分块"}

        self._ensure_vectorstore()

        doc_id = str(uuid.uuid4())
        chunk_ids = [f"{doc_id}_chunk_{i}" for i in range(len(chunks))]

        self._vectorstore.add_texts(
            texts=chunks,
            metadatas=[
                {"doc_id": doc_id, "filename": src.name, "chunk": i}
                for i in range(len(chunks))
            ],
            ids=chunk_ids,
        )

        info = {
            "id": doc_id,
            "filename": src.name,
            "chunks": len(chunks),
            "created_at": datetime.now(timezone.utc).isoformat(),
        }

        index = self._load_index()
        index.append(info)
        self._save_index(index)

        return info

    def search(self, query: str, k: int = 5) -> list[Any]:
        """检索最相关的 k 个文本块。"""
        self._ensure_vectorstore()
        try:
            return self._vectorstore.similarity_search(query, k=k)
        except Exception:
            return []

    def list_documents(self) -> list[dict[str, Any]]:
        """列出已索引的文档列表。"""
        return self._load_index()

    def delete_document(self, doc_id: str) -> bool:
        """从索引中删除文档。"""
        self._ensure_vectorstore()

        index = self._load_index()
        matched = [d for d in index if d["id"] == doc_id]
        if not matched:
            return False

        doc_info = matched[0]
        num_chunks = doc_info["chunks"]

        try:
            self._vectorstore.delete(
                ids=[f"{doc_id}_chunk_{i}" for i in range(num_chunks)]
            )
        except Exception:
            pass  # Chroma 删除失败不阻塞索引清理

        index = [d for d in index if d["id"] != doc_id]
        self._save_index(index)
        return True

    # ── 持久化 ────────────────────────────────────────────────────

    def _load_index(self) -> list[dict[str, Any]]:
        if self._index_path.exists():
            with open(self._index_path, "r", encoding="utf-8") as f:
                return json.load(f)
        return []

    def _save_index(self, index: list[dict[str, Any]]):
        with open(self._index_path, "w", encoding="utf-8") as f:
            json.dump(index, f, ensure_ascii=False, indent=2)
